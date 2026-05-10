import re
import json
import unicodedata
from pathlib import Path
from tkinter import Tk, filedialog, messagebox

from docx import Document

# =========================
# Seleccionar archivo
# =========================

Tk().withdraw()

ruta_archivo = filedialog.askopenfilename(
    title="Selecciona la plantilla Word",
    filetypes=[("Documentos Word", "*.docx")]
)

if not ruta_archivo:
    # print("No se seleccionó ningún archivo.")
    # input("Presiona Enter para salir...")
    exit()

ruta_archivo = Path(ruta_archivo)

# =========================
# Nombres de salida
# =========================

archivo_salida = ruta_archivo.with_name(
    # ruta_archivo.stem + "_jinja.docx"
    "plantilla_jinja.docx"
)

archivo_json = ruta_archivo.with_name(
    # ruta_archivo.stem + "_variables.json"
    "plantilla_variables.json"
)

# =========================
# Abrir documento
# =========================

doc = Document(str(ruta_archivo))

# Diccionario de variables
mapa_variables = {}

# =========================
# Funciones
# =========================

def quitar_acentos(texto):
    texto = unicodedata.normalize('NFD', texto)

    return ''.join(
        c for c in texto
        if unicodedata.category(c) != 'Mn'
    )

def convertir_variable(texto_original):
    
    print("input:", texto_original)
    
    # Minúsculas
    texto = texto_original.lower()
    
    # Quitar acentos
    texto = quitar_acentos(texto)

    # Eliminar caracteres especiales
    texto = re.sub(r"[^a-z0-9_ ]", "", texto)

    # Eliminar espacios al inicio y al final
    texto = texto.strip()

    # Ignorar vacío
    if not texto:
        return f"[{texto_original}]"

    # Debe comenzar con letra
    if not re.match(r"^[a-z]", texto):
        return f"[{texto_original}]"

    # Espacios → _
    texto = texto.replace(" ", "_")

    # Guardar mapeo
    if texto not in mapa_variables:
        mapa_variables[texto] = texto_original.strip()

    return f"{{{{ {texto} }}}}"

# Buscar [VARIABLE]
patron = re.compile(r"\[([^\[\]]+)\]")

# =========================
# Procesar párrafos
# =========================

def procesar_parrafos(paragraphs):

    for paragraph in paragraphs:

        full_text = ""
        char_map = []

        # -------------------------------------------------
        # Construir texto completo + mapa de caracteres
        # -------------------------------------------------

        for run_index, run in enumerate(paragraph.runs):

            for char_index, char in enumerate(run.text):

                char_map.append({
                    "run": run,
                    "run_index": run_index,
                    "char_index": char_index
                })

                full_text += char

        # -------------------------------------------------
        # Buscar variables
        # -------------------------------------------------

        matches = list(patron.finditer(full_text))

        # ---------------------------------------------
        # Registrar variables en orden natural
        # ---------------------------------------------

        for match in matches:

            variable_original = match.group(1)

            convertir_variable(variable_original)

        # Procesar en reversa para no romper índices
        for match in reversed(matches):

            variable_original = match.group(1)

            nueva_variable = convertir_variable(
                variable_original
            )

            start = match.start()
            end = match.end() - 1

            # ---------------------------------------------
            # Obtener información de inicio y fin
            # ---------------------------------------------

            start_info = char_map[start]
            end_info = char_map[end]

            start_run = start_info["run"]
            end_run = end_info["run"]

            start_char = start_info["char_index"]
            end_char = end_info["char_index"]

            # ---------------------------------------------
            # Caso 1:
            # Todo está en un mismo run
            # ---------------------------------------------

            if start_run == end_run:

                texto = start_run.text

                nuevo_texto = (
                    texto[:start_char] +
                    nueva_variable +
                    texto[end_char + 1:]
                )

                start_run.text = nuevo_texto

            # ---------------------------------------------
            # Caso 2:
            # Variable repartida en varios runs
            # ---------------------------------------------

            else:

                # Texto antes de la variable
                texto_inicio = start_run.text[:start_char]

                # Texto después de la variable
                texto_final = end_run.text[end_char + 1:]

                # Reemplazar en primer run
                start_run.text = (
                    texto_inicio +
                    nueva_variable
                )

                # Vaciar runs intermedios
                for i in range(
                    start_info["run_index"] + 1,
                    end_info["run_index"]
                ):

                    paragraph.runs[i].text = ""

                # Conservar sobrante del último run
                end_run.text = texto_final

procesar_parrafos(doc.paragraphs)

# =========================
# Procesar tablas
# =========================

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:

            procesar_parrafos(cell.paragraphs)

# =========================
# Guardar archivos
# =========================

doc.save(str(archivo_salida))

with open(archivo_json, "w", encoding="utf-8") as f:
    json.dump(
        mapa_variables,
        f,
        ensure_ascii=False,
        indent=4
    )

# =========================
# Final
# =========================

# print("Proceso completado correctamente.\n")

# print(f"Documento generado:")
# print(f"  {archivo_salida}")

# print(f"\nMapa JSON generado:")
# print(f"  {archivo_json}")

# input("\nPresiona Enter para salir...")

messagebox.showinfo(
    "Proceso completado",
    f"Documento generado:\n{archivo_salida}\n\n"
    f"Mapa JSON generado:\n{archivo_json}"
)